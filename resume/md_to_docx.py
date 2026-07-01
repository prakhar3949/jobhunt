"""Convert the resume Markdown files in this folder to formatted .docx.

Purpose-built for the structure used in the resume .md files:
  # Name                -> title (centered, large)
  **bold line**         -> subtitle/role (centered)
  contact line          -> centered small line (the line right after subtitle, before first ---)
  ## Section            -> section heading + bottom border
  ### Job header        -> bold left + right-aligned date (split on '  ·  ')
  *italic*              -> italic stack line
  - bullet              -> bullet, with **bold** inline runs
  ---                   -> ignored (visual separator handled by section borders)
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).resolve().parent
DARK = RGBColor(0x1a, 0x1a, 0x2e)
ACCENT = RGBColor(0x2a, 0x4d, 0x7a)
GREY = RGBColor(0x55, 0x55, 0x55)


def add_bottom_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2a4d7a')
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_runs_with_bold(paragraph, text, size=10, color=None):
    """Render text containing **bold** spans into runs."""
    for i, part in enumerate(re.split(r'\*\*(.+?)\*\*', text)):
        if part == '':
            continue
        run = paragraph.add_run(part)
        run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = color
        if i % 2 == 1:  # odd indices are the captured bold groups
            run.bold = True


def strip_links(text):
    # [label](url) -> label
    return re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)


def convert(md_path: Path):
    lines = md_path.read_text(encoding='utf-8').splitlines()
    doc = Document()
    # base style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(36)
        section.left_margin = section.right_margin = Pt(54)

    state = 'title'  # title -> subtitle -> contact -> body
    i = 0
    n = len(lines)
    while i < n:
        raw = lines[i]
        line = raw.strip()
        i += 1
        if not line:
            continue

        if line.startswith('# ') and state == 'title':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line[2:].strip())
            r.bold = True
            r.font.size = Pt(22)
            r.font.color.rgb = DARK
            state = 'subtitle'
            continue

        if state == 'subtitle' and line.startswith('**'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line.strip('*').strip())
            r.bold = True
            r.font.size = Pt(11.5)
            r.font.color.rgb = ACCENT
            state = 'contact'
            continue

        if state == 'contact' and not line.startswith('#') and line != '---':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(strip_links(line))
            r.font.size = Pt(9)
            r.font.color.rgb = GREY
            state = 'body'
            continue

        if line == '---':
            continue

        if line.startswith('## '):
            p = doc.add_paragraph()
            p.space_before = Pt(8)
            r = p.add_run(line[3:].strip().upper())
            r.bold = True
            r.font.size = Pt(12)
            r.font.color.rgb = ACCENT
            add_bottom_border(p)
            state = 'body'
            continue

        if line.startswith('### '):
            content = line[4:].strip()
            parts = content.split('  ·  ')
            title = parts[0]
            date = parts[-1] if len(parts) > 1 else ''
            p = doc.add_paragraph()
            p.space_before = Pt(6)
            # tab stop at right margin for the date
            usable = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
            p.paragraph_format.tab_stops.add_tab_stop(usable, WD_TAB_ALIGNMENT.RIGHT)
            add_runs_with_bold(p, '**' + title + '**', size=11, color=DARK)
            if date:
                r = p.add_run('\t' + date)
                r.font.size = Pt(9.5)
                r.font.color.rgb = GREY
                r.italic = True
            state = 'body'
            continue

        if line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            add_runs_with_bold(p, strip_links(line[2:].strip()), size=10)
            continue

        if line.startswith('*') and line.endswith('*') and not line.startswith('**'):
            p = doc.add_paragraph()
            r = p.add_run(line.strip('*'))
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = GREY
            continue

        # default paragraph (e.g. summary text, possibly with **bold**)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        add_runs_with_bold(p, strip_links(line), size=10)

    out = md_path.with_suffix('.docx')
    doc.save(out)
    print('wrote', out.name)


if __name__ == '__main__':
    for md in sorted(HERE.glob('Prakhar-Goyal-Resume-*.md')):
        convert(md)
